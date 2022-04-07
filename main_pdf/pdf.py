from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx2pdf import convert
from pathlib import Path
import time

from py_tarif.vars import all_control_wifi_dop, all_wifi_guest_dop, wats_tariff_cat, wats_tariff, internet_tariff

from py_tarif.tarif import data as department_tariff

wifi_guest_tp_names_pdf = {
    'start': 'ТП Начальный',
    'improve': 'ТП Улушенный',
    'extended': 'ТП Расширенный'
}

wats_tp_to_name_pdf = {
    'bus/st': 'Бизнес. Стандарт',
    'bus/unlim': 'Бизнес. Безлимит',
    'virt/st': 'Виртуальный номер. Стандарт',
    'virt/unlim': 'Виртуальный номер. Безлимит',
    'bus/korp20': 'Бизнес. Корпорация 20',
    'bus/korp40': 'Бизнес. Корпорация 40',

    '3r/0m-500': '3 рабочих места, 0 минут',
    '5r/600m-1100': '5 рабочих мест, 600 минут',
    '10r/1200m-1900': '10 рабочих мест, 1200 минут',
    '15r/3000m-4900': '15 рабочих мест, 3000 минут',
    '30r/5000m-7900': '30 рабочих мест, 5000 минут'
}

wats_num_abon_to_name_pdf = [
        '1user',
        '23user',
        'do5',
        'more5user',
        '20user',
        'more20User',
        '40user',
        'more40User'
    ]

vn_tp_to_name_pdf = {
    'vnStan': '"ВН" Стандартный',
    'vnBase': '"ВН" Базовый 2.0',
    'stan0': 'Стандартный',
    'stan10': 'Стандартный-10',
    'stan20': 'Стандартный-20',
    'base20': 'Базовый 2.0',
    'base20mlm': 'Базовый 2.0 mlm',
    'base20dig': 'Базовый 2.0_Digital',
    'base2010': 'Базовый 2.0 - 10',
    'base2010mlm': 'Базовый 2.0 - 10 mlm',
    'base2010dig': 'Базовый 2.0_Digital - 10',
    'base2020': 'Базовый 2.0 - 20',
    'base2020mlm': 'Базовый 2.0 - 20 mlm',
    'base2020dig': 'Базовый 2.0_Digital - 20',
    'base2030': 'Базовый 2.0 - 30',
    'base2030mlm': 'Базовый 2.0 - 30 mlm',
    'base2030dig': 'Базовый 2.0_Digital - 30',
}

vn_save_to_name_pdf = {
    'reg': 'Регистратор',
    'complete': 'Полная запись',
    'actions': 'Запись только событий',
    'econom': 'Режим экономии трафика'
}

iptv_hot_main_ch_pdf = {
    '1mainChNoInt': 'ОПТ без интеграции',
    '2mainChPMS': 'ОПТ при интеграции (PMS)',
    '3mainChDVB': 'ОПТ при подключении DVB-C',
    '4mainChSmart': 'ОПТ при подключении на SmartTV',
    '5mainChSmartPMS': 'ОПТ на SmartTV при интеграции (PMS)',

    '1st': 'Стандарт',
    '1lux': 'Люкс',
    '1prestige': 'Престиж',
    '1st/daily': 'Стандарт (суточный)',
    '1lux/daily': 'Люкс (суточный)',
    '1prestige/daily': 'Престиж (суточный)',

    '2st/int': 'Стандарт (с интеграцией)',
    '2lux/int': 'Люкс (с интеграцией)',
    '2prestige/int': 'Престиж (с интеграцией)',
    '2st/daily/int': 'Стандарт (суточный, с интеграцией)',
    '2lux/daily/int': 'Люкс (суточный, с интеграцией)',
    '2prestige/daily/int': 'Престиж (суточный, с интеграцией)',

    '3st/dvb': 'Стандарт DVB-C',
    '3lux/dvb': 'Люкс DVB-C',
    '3prestige/dvb': 'Престиж DVB-C',
    '3st/light/dvb': 'Стандарт Лайт DVB-C',

    '4st/smart': 'Стандарт SmartTV',
    '4lux/smart': 'Люкс SmartTV',
    '4prestige/smart': 'Престиж SmartTV',
    '4st/daily/smart': 'Стандарт SmartTV (суточный)',
    '4lux/daily/smart': 'Люкс SmartTV (суточный)',
    '4prestige/daily/smart': 'Престиж SmartTV (суточный)',

    '5st/smart/int': 'Стандарт SmartTV (с интеграцией)',
    '5lux/smart/int': 'Люкс SmartTV (с интеграцией)',
    '5prestige/smart/int': 'Престиж SmartTV (с интеграцией)',
    '5st/daily/smart/int': 'Стандарт SmartTV (суточный, с интеграцией)',
    '5lux/daily/smart/int': 'Люкс SmartTV (суточный, с интеграцией)',
    '5prestige/daily/smart/int': 'Престиж SmartTV (суточный, с интеграцией)',
    '5partner': 'Пакет "Партнерский стандарт"',
    '5partner/int': 'Пакет "Партнерский стандарт" (с интеграцией)'
}

iptv_hot_choices_pdf = {
    '1/10': '1-10 пакетов: ',
    '11/20': '11-20 пакетов: ',
    '21/30': '21-30 пакетов: ',
    '30+': 'свыше 30 пакетов: ',

    '1-0': '1 канал: 0 ₽ ежемес.',
    '2-60': '2 канала: 60 ₽ ежемес.',
    '3-75': '3 канала: 75 ₽ ежемес.',
    '4-95': '4 канала: 95 ₽ ежемес.',
    '5-120': '5 каналов: 120 ₽ ежемес.'
}

iptv_hot_dop_ch_names_pdf = {
    'ihkhlprime': 'Телеканал KHL Prime (149 ₽)',
    'ihmatch': 'Матч! (100 ₽)',
    'ihsupermatch': 'Суперматч! (450 ₽)',
    'ihadult': 'Пакет для взрослых (500 ₽)',
    'ihadult+': 'Пакет для взрослых + (1200 ₽)',
    'ihadultSut': 'Пакет для взрослых (суточный) (120 ₽)',
    'ihadult+Sut': 'Пакет для взрослых + (суточный) (360 ₽)',
    'ihchina': 'Пакет Китай (150 ₽)',
    'ihchinaSut': 'Пакет Китай (суточный) (13 ₽)',
    'ihpass/dopchannel': 'Пропустить!'
}

ip_service_names_pdf = {
    'public': 'Публичный показ',
    'private': 'Непубличный показ',

    'sfera/other': 'Для сферы услуг, на интернете др. провайдера',
    'sfera+/other': 'Для сферы услуг+, на интернете др. провайдера',
    'dvb/sfera/nodop': 'Для сферы услуг DVB-C (без доп.)',
    'sfera': 'Для сферы услуг',
    'sfera+': 'Для сферы услуг +',
    'kids': 'Для детей',
    'kids+': 'Для детей +',
    'dvb/lite': 'DVB-C - Публичный лайт',
    'dvb/sfera': 'Для сферы услуг DVB-C',
    'dvb/sfera+': 'Для сферы услуг+ DVB-C',
    'dvb/kids': 'Для детей DVB-C',
    'dvb/kids+': 'Для детей + DVB-C',

    'office': 'ТВ для офиса',
    'office+': 'ТВ для офиса +',
    'office/sm': 'ТВ для офиса SMART TV',
    'office+/sm': 'ТВ для офиса + SMART TV',
    'office/other/sm': 'ТВ для офиса на интернете др. провайдера SMART',
    'office+/other/sm': 'ТВ для офиса + на интернете др. провайдера SMART',
    'office/other': 'ТВ для офиса на интернете др.провайдера',
    'office+/other': 'ТВ для офиса + на интернете др.провайдера',
    'dvb/office/nodop': 'ТВ для офиса DVB-C'
}

ip_dop_ch_names_pdf = {
    'khlprime': 'Телеканал KHL Prime',
    'match': 'Матч!',
    'supermatch': 'Суперматч!',
    'china': 'Пакет Китай',
    'smart': 'SMART TV публичный показ',
    'adult+': 'Пакет телеканалов «Пакет для взрослых +»',
    'sm.kids': 'Для детей',
    'sm.kids+': 'Для детей +',
    'sm.sfera': 'Для сферы услуг',
    'sm.sfera+': 'Для сферы услуг +',
    'sm.sfera/other': 'Для сферы услуг; др. провайдер',
    'sm.sfera+/other': 'Для сферы услуг +; др. провайдер'
}

ip_prokat_names_pdf = {
    1: 'Категория 1 (100 ₽ единовременно)',
    2: 'Категория 2 (150 ₽ единовременно)',
    3: 'Категория 3 (200 ₽ единовременно)',
    4: 'Категория 4 (250 ₽ единовременно)',
    5: 'Категория 5 (300 ₽ единовременно)',
    6: 'Категория 6 (350 ₽ единовременно)',
    7: 'Категория 7 (400 ₽ единовременно)',
    8: 'Категория 8 (450 ₽ единовременно)',
    9: 'Категория 9 (500 ₽ единовременно)',
    10: 'Категория 10 (750 ₽ единовременно)',
    11: 'Категория 11 (1000 ₽ единовременно)'
}

internet_tariffs_names_pdf = {
    'delovoj': 'Деловой',
    "uskorenie": "Ускорение",
    "ozhn": "ОЖН",
    "skorost": "Скорость",
    "m2mkontrol": "М2М Контроль"
}

# tariffs_names = {
#     'delovoj': 'Деловой',
#     "bytvpljuse": "Быть в плюсе",
#     "uskorenie": "Ускорение",
#     "ozhn": "ОЖН",
#     "skorost": "Скорость",
#     "m2mkontrol": "М2М Контроль"
# }

async def tostr_(lis):
    P = -3
    C = ' '
    lis.insert(P, C)
    ret = ''
    for i in lis:
        ret += i
    return ret

async def getprice(item):
    for specs in internet_tariff[item.split('_')[1]][item.split('_')[2]][item.split('_')[3]][item.split('_')[5]]:
        if specs['speed'] == item.split('_')[4]:
            return specs["price"]

async def create_pdf(uid, tar_list, tar_order):

    for tar_id in range(len(tar_order[:-1])):

        curdir = Path.cwd() / 'main_pdf'

        if tar_order[tar_id] == 'wifi/control':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                #sms
                sms_cell = 1 if 'nosms' in tar_list[tar_id][i].split('_')[3] else 0
                empty_sms_cell = 0 if sms_cell == 1 else 1
                sms_price = tar_list[tar_id][i].split('_')[3].split('-')[1]

                doc.tables[1].cell(empty_sms_cell + 1, i + 1).text = str('--')
                doc.tables[1].cell(empty_sms_cell + 1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(empty_sms_cell + 1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(empty_sms_cell + 1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                doc.tables[1].cell(sms_cell + 1, i + 1).text = str(sms_price + ' ₽ / месяц')
                doc.tables[1].cell(sms_cell + 1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(sms_cell + 1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(sms_cell + 1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                #wifi
                org = tar_list[tar_id][i].split('_')[2].split('-')[1]

                doc.tables[1].cell(3, i + 1).text = str(org + ' ₽ за устройство')
                doc.tables[1].cell(3, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.name = 'Arial'


                #filter
                filter_price = '--' if 'filter' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'filter' in a][0]) + ' ₽'

                doc.tables[1].cell(4, i + 1).text = str(filter_price )
                doc.tables[1].cell(4, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.name = 'Arial'


                # block
                block_price = '--' if 'block' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'block' in a][0]) + ' ₽'

                doc.tables[1].cell(5, i + 1).text = str(block_price)
                doc.tables[1].cell(5, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # radar
                try:
                    radar = [a for a in tar_list[tar_id][i].split('_') if 'radar' in a][0]
                    radar_pos = int(radar.split('-')[-1])
                    radar_device = radar.split('-')[-2]
                    radar_price_device = radar_device.split('/')[1]
                    radar_months = radar_device.split('/')[0]
                    if radar_months == 'once':
                        radar_months = 'Срок: единовременно'
                    else:
                        radar_months = str('Срок: ') + radar_months.replace('m', ' месяцев')
                    radar_price_month = all_control_wifi_dop['dop']['radar'][radar_device][str(radar_pos)]
                    radar_str = str(radar_price_device) + ' ₽ за устройство; \n' + str(radar_months) + '\n' + str(radar_price_month) + ' ₽ за лицензию'
                except:
                    radar_str = '--'

                radar_list = [1, 3, 5, 10]
                for j in range(len(radar_list)):
                    try:
                        radar_pos = radar_pos
                    except:
                        radar_pos = 11
                    # print(j)
                    if radar_list[j] == radar_pos:
                        doc.tables[1].cell(6+j, i + 1).text = str(radar_str)
                        doc.tables[1].cell(6+j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                        doc.tables[1].cell(6+j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                        doc.tables[1].cell(6+j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'
                    else:
                        doc.tables[1].cell(6 + j, i + 1).text = str('--')
                        doc.tables[1].cell(6 + j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                        doc.tables[1].cell(6 + j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                        doc.tables[1].cell(6 + j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        elif tar_order[tar_id] == 'wifi/guest':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                #naimenovanie
                wg_name = wifi_guest_tp_names_pdf[tar_list[tar_id][i].split('_')[3].split('-')[0]]

                doc.tables[1].cell(1, i + 1).text = str(wg_name)
                doc.tables[1].cell(1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # organizacia
                wg_org = tar_list[tar_id][i].split('_')[3].split('-')[1]
                wg_text = str(wg_org.split('/')[0]) + ' ₽ за устройство \n' + str(wg_org.split('/')[1]) + ' ₽ за точку (в мес.)'

                doc.tables[1].cell(2, i + 1).text = str(wg_text)
                doc.tables[1].cell(2, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                #speed
                wg_spd_price = tar_list[tar_id][i].split('_')[4].split('-')[1]
                wg_spd_val = '0'
                for k, v in all_wifi_guest_dop['speed'][tar_list[tar_id][i].split('_')[1]].items():
                    if v == int(wg_spd_price):
                        wg_spd_val = k
                wg_spd_txt = str(wg_spd_val) + " Мбит/с; " + str(wg_spd_price) + " ₽ / месяц"
                doc.tables[1].cell(3, i + 1).text = str(wg_spd_txt)
                doc.tables[1].cell(3, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # filter
                filter_price = '--' if 'filter' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'filter' in a][0]) + ' ₽'

                doc.tables[1].cell(4, i + 1).text = str(filter_price)
                doc.tables[1].cell(4, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # block
                block_price = '--' if 'block' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'block' in a][0]) + ' ₽'

                doc.tables[1].cell(5, i + 1).text = str(block_price)
                doc.tables[1].cell(5, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # adv

                for j in range(len(['access', '12m', 'extended'])):
                    adv_wg_price = '--'
                    if 'adv-' + ['access', '12m', 'extended'][j] in tar_list[tar_id][i]:
                        adv_wg_price = str(all_wifi_guest_dop['dop']['adv'][['access', '12m', 'extended'][j]]) + ' ₽'

                    doc.tables[1].cell(6 + j, i + 1).text = str(adv_wg_price)
                    doc.tables[1].cell(6 + j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                    doc.tables[1].cell(6 + j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                    doc.tables[1].cell(6 + j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # radar
                try:
                    radar = [a for a in tar_list[tar_id][i].split('_') if 'radar' in a][0]
                    radar_pos = int(radar.split('-')[-1])
                    radar_device = radar.split('-')[-2]
                    radar_price_device = radar_device.split('/')[1]
                    radar_months = radar_device.split('/')[0]
                    if radar_months == 'once':
                        radar_months = 'Срок: единовременно'
                    else:
                        radar_months = str('Срок: ') + radar_months.replace('m', ' месяцев')
                    radar_price_month = all_wifi_guest_dop['dop']['radar'][radar_device][str(radar_pos)]
                    radar_str = str(radar_price_device) + ' ₽ за устройство; \n' + str(radar_months) + '\n' + str(radar_price_month) + ' ₽ за лицензию'
                except:
                    radar_str = '--'

                radar_list = [1, 3, 5, 10]
                for j in range(len(radar_list)):
                    try:
                        radar_pos = radar_pos
                    except:
                        radar_pos = 11
                    if radar_list[j] == radar_pos:
                        doc.tables[1].cell(9 + j, i + 1).text = str(radar_str)
                        doc.tables[1].cell(9 + j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                        doc.tables[1].cell(9 + j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                        doc.tables[1].cell(9 + j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'
                    else:
                        doc.tables[1].cell(9 + j, i + 1).text = str('--')
                        doc.tables[1].cell(9 + j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                        doc.tables[1].cell(9 + j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                        doc.tables[1].cell(9 + j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        elif tar_order[tar_id] == 'wats/busy':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                #naimenovanie
                wab_name = wats_tp_to_name_pdf[tar_list[tar_id][i].split('_')[3]]

                doc.tables[1].cell(1, i + 1).text = str(wab_name)
                doc.tables[1].cell(1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # number_cat

                for j in range(len(['platinum-100000', 'gold-40000', 'silver-15000', 'bronze-5000', 'nocat'])):
                    cat_wab_price = '--'
                    if ['platinum-100000', 'gold-40000', 'silver-15000', 'bronze-5000', 'nocat'][j] in tar_list[tar_id][i]:
                        cat_wab_price = str(wats_tariff_cat[['platinum-100000', 'gold-40000', 'silver-15000', 'bronze-5000', 'nocat'][j]]) + ' ₽'
                        if cat_wab_price == 'in/abon ₽':
                            cat_wab_price = 'Платеж за 1 номер включен в абон. плату услуги Виртуальная АТС'

                    doc.tables[1].cell(2 + j, i + 1).text = str(cat_wab_price)
                    doc.tables[1].cell(2 + j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                    doc.tables[1].cell(2 + j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                    doc.tables[1].cell(2 + j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # users

                for j in range(len(wats_num_abon_to_name_pdf)):
                    us_price_wab = '--' if wats_num_abon_to_name_pdf[j] not in tar_list[tar_id][i] else \
                        str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if wats_num_abon_to_name_pdf[j] in a][0]) + ' ₽'
                    doc.tables[1].cell(7 + j, i + 1).text = str(us_price_wab)
                    doc.tables[1].cell(7 + j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                    doc.tables[1].cell(7 + j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                    doc.tables[1].cell(7 + j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # vid_back
                vid_back = '--' if 'vid/back' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'vid/back' in a][0]) + ' ₽'

                doc.tables[1].cell(15, i + 1).text = str(vid_back)
                doc.tables[1].cell(15, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(15, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(15, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # virt_contact 2 place
                virt_cont_2 = '--' if 'virt/contact/2' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'virt/contact/2' in a][0]) + ' ₽'

                doc.tables[1].cell(16, i + 1).text = str(virt_cont_2)
                doc.tables[1].cell(16, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(16, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(16, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # virt_contact dop place
                virt_cont_dop = '--' if 'virt/contact/dop' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'virt/contact/dop' in a][0]) + ' ₽'

                doc.tables[1].cell(17, i + 1).text = str(virt_cont_dop)
                doc.tables[1].cell(17, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(17, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(17, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # integration other crm
                other_crm = '--' if 'other/crm' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'other/crm' in a][0]) + ' ₽'

                doc.tables[1].cell(18, i + 1).text = str(other_crm)
                doc.tables[1].cell(18, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(18, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(18, i + 1).paragraphs[0].runs[0].font.name = 'Arial'


            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        elif tar_order[tar_id] == 'wats/packmin':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                #naimenovanie

                wap_name = wats_tp_to_name_pdf[tar_list[tar_id][i].split('_')[3]]

                doc.tables[1].cell(1, i + 1).text = str(wap_name)
                doc.tables[1].cell(1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # mesta
                doc.tables[1].cell(2, i + 1).text = str(wap_name.split()[0])
                doc.tables[1].cell(2, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # minuti
                doc.tables[1].cell(3, i + 1).text = str(wap_name.split()[3]) + ' минут'
                doc.tables[1].cell(3, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # categoria
                doc.tables[1].cell(4, i + 1).text = str('Включен городской номер')
                doc.tables[1].cell(4, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # price
                doc.tables[1].cell(5, i + 1).text = str(tar_list[tar_id][i].split('_')[3].split('-')[1]) + ' ₽'
                doc.tables[1].cell(5, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # dop_workp - оставил чтобы всегда было
                wap_dop_place = wats_tariff['pack/min']['dop']['more/workp'][tar_list[tar_id][i].split('_')[3]]
                doc.tables[1].cell(6, i + 1).text = str(wap_dop_place)
                doc.tables[1].cell(6, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(6, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(6, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # mobile_workp
                mobile_workp = '--' if 'mobile/workp' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'mobile/workp' in a][0]) + ' ₽'

                doc.tables[1].cell(7, i + 1).text = str(mobile_workp)
                doc.tables[1].cell(7, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(7, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(7, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # vid_back
                vid_back = '--' if 'vid/back' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'vid/back' in a][0]) + ' ₽'

                doc.tables[1].cell(8, i + 1).text = str(vid_back)
                doc.tables[1].cell(8, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(8, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(8, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # anal
                anal = '--' if 'anal' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'anal' in a][0]) + ' ₽'

                doc.tables[1].cell(9, i + 1).text = str(anal)
                doc.tables[1].cell(9, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(9, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(9, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # virt_contact 2 place
                virt_cont_2 = '--' if 'virt/contact/2' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'virt/contact/2' in a][0]) + ' ₽'

                doc.tables[1].cell(10, i + 1).text = str(virt_cont_2)
                doc.tables[1].cell(10, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(10, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(10, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # virt_contact dop place
                virt_cont_dop = '--' if 'virt/contact/dop' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'virt/contact/dop' in a][0]) + ' ₽'

                doc.tables[1].cell(11, i + 1).text = str(virt_cont_dop)
                doc.tables[1].cell(11, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(11, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(11, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # integration amocrm
                amocrm = '--' if 'amocrm' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'amocrm' in a][0]) + ' ₽'

                doc.tables[1].cell(12, i + 1).text = str(amocrm)
                doc.tables[1].cell(12, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(12, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(12, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # integration api crm
                api_crm = '--' if 'other/crm' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'other/crm' in a][0]) + ' ₽'

                doc.tables[1].cell(13, i + 1).text = str(api_crm)
                doc.tables[1].cell(13, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(13, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(13, i + 1).paragraphs[0].runs[0].font.name = 'Arial'


            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        elif tar_order[tar_id] == 'vn':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                #naimenovanie
                vn_name = vn_tp_to_name_pdf[tar_list[tar_id][i].split('_')[2]]

                doc.tables[1].cell(1, i + 1).text = str(vn_name)
                doc.tables[1].cell(1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # tp
                vn_tp = vn_tp_to_name_pdf[tar_list[tar_id][i].split('_')[3]]

                doc.tables[1].cell(2, i + 1).text = str(vn_tp)
                doc.tables[1].cell(2, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # sutki hranenia
                sut_vn = tar_list[tar_id][i].split('_')[4]

                doc.tables[1].cell(3, i + 1).text = str(sut_vn) + ' суток'
                doc.tables[1].cell(3, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # vid zapisi
                try:
                    vid_rec_vn_str = vn_save_to_name_pdf[tar_list[tar_id][i].split('_')[5]]
                except:
                    vid_rec_vn = vn_save_to_name_pdf[tar_list[tar_id][i].split('_')[5].split('-')[0]]
                    vid_rec_vn_str = vid_rec_vn + f" {tar_list[tar_id][i].split('_')[5].split('-')[1]} ₽"

                doc.tables[1].cell(4, i + 1).text = str(vid_rec_vn_str)
                doc.tables[1].cell(4, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # traffic

                if '₽' in vid_rec_vn_str:
                    tr_vn_str = '--'
                else:
                    traffic_vn = tar_list[tar_id][i].split('_')[6]
                    tr_spd = traffic_vn.split('-')[0]
                    tr_price = traffic_vn.split('-')[1]
                    tr_vn_str = str(tr_spd) + f'{" Мбит/с" if float(tr_spd) < 150 else " Кбит/с"}' + f": {str(tr_price)} ₽"

                doc.tables[1].cell(5, i + 1).text = str(tr_vn_str)
                doc.tables[1].cell(5, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # sms notice
                for j in range(len(['100/30', '500/90', '1000/180', '5000/365'])):
                    if ['100/30', '500/90', '1000/180', '5000/365'][j] in tar_list[tar_id][i]:
                        sms_notice_vn_str = str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if any(z in a for z in ['100/30', '500/90', '1000/180', '5000/365'])][0]) + ' ₽'
                    else:
                        sms_notice_vn_str = '--'
                    doc.tables[1].cell(6+j, i + 1).text = str(sms_notice_vn_str)
                    doc.tables[1].cell(6+j, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                    doc.tables[1].cell(6+j, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                    doc.tables[1].cell(6+j, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # count visit

                vis_c_vn = '--' if 'visCount' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'visCount' in a][0]) + ' ₽'

                doc.tables[1].cell(10, i + 1).text = str(vis_c_vn)
                doc.tables[1].cell(10, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(10, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(10, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # q lenght

                lenq_vn = '--' if 'lenQ' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'lenQ' in a][0]) + ' ₽'

                doc.tables[1].cell(11, i + 1).text = str(lenq_vn)
                doc.tables[1].cell(11, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(11, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(11, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # facerec

                facerec_vn = '--' if 'faceRec' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'faceRec' in a][0]) + ' ₽'

                doc.tables[1].cell(12, i + 1).text = str(facerec_vn)
                doc.tables[1].cell(12, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(12, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(12, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # controlEmp

                control_emp_vn = '--' if 'controlEmp' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'controlEmp' in a][0]) + ' ₽'

                doc.tables[1].cell(13, i + 1).text = str(control_emp_vn)
                doc.tables[1].cell(13, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(13, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(13, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # heatMap

                heatMap_vn = '--' if 'heatMap' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'heatMap' in a][0]) + ' ₽'

                doc.tables[1].cell(14, i + 1).text = str(heatMap_vn)
                doc.tables[1].cell(14, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(14, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(14, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # plate_rec_vn

                plate_rec_vn = '--' if 'autoNom' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'autoNom' in a][0]) + ' ₽'

                doc.tables[1].cell(15, i + 1).text = str(plate_rec_vn)
                doc.tables[1].cell(15, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(15, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(15, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # left things

                left_vn = '--' if 'leftS' not in tar_list[tar_id][i] else \
                    str([a.split('-')[1] for a in tar_list[tar_id][i].split('_') if 'leftS' in a][0]) + ' ₽'

                doc.tables[1].cell(16, i + 1).text = str(left_vn)
                doc.tables[1].cell(16, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(16, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(16, i + 1).paragraphs[0].runs[0].font.name = 'Arial'


            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        elif tar_order[tar_id] == 'iptv/hot':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                #naimenovanie
                iph_name_1 = iptv_hot_main_ch_pdf[tar_list[tar_id][i].split('_')[2]].replace('ОПТ', 'Основные пакеты ТК') + '\n'
                iph_name_2 = 'Пакет: ' + iptv_hot_main_ch_pdf[tar_list[tar_id][i].split('_')[3].split('-')[0]]
                iph_name = iph_name_1 + iph_name_2

                doc.tables[1].cell(1, i + 1).text = str(iph_name)
                doc.tables[1].cell(1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                #num_pack and price
                if tar_list[tar_id][i].split('_')[3].split('-')[0] == tar_list[tar_id][i].split('_')[3]:
                    iph_num = iptv_hot_choices_pdf[tar_list[tar_id][i].split('_')[4].split('-')[0]].replace(': ', '')
                    iph_price = tar_list[tar_id][i].split('_')[4].split('-')[1] + ' ₽'
                else:
                    iph_num = '1 пакет'
                    iph_price = tar_list[tar_id][i].split('_')[3].split('-')[1] + ' ₽'

                doc.tables[1].cell(2, i + 1).text = str(iph_num)
                doc.tables[1].cell(2, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                doc.tables[1].cell(3, i + 1).text = str(iph_price)
                doc.tables[1].cell(3, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                #dop_ch
                ih_dop = str()
                for j in tar_list[tar_id][i].split('_'):
                    try:
                        ih_dop += iptv_hot_dop_ch_names_pdf[j.split('-')[0]] + ' \n'
                    except:
                        pass

                if len(ih_dop) == 0:
                    ih_dop = '--'

                doc.tables[1].cell(4, i + 1).text = str(ih_dop)
                doc.tables[1].cell(4, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # infoch
                ih_info = '--' if not any(j in tar_list[tar_id][i] for j in ['1-0', '2-60', '3-75', '4-95', '5-120']) else \
                    str([iptv_hot_choices_pdf[a] for a in tar_list[tar_id][i].split('_') if
                         any(j in a for j in ['1-0', '2-60', '3-75', '4-95', '5-120'])][0])

                doc.tables[1].cell(5, i + 1).text = str(ih_info)
                doc.tables[1].cell(5, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.name = 'Arial'


            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        elif tar_order[tar_id] == 'iptv':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                # iptv naimenovanie
                ip_name_1 = ip_service_names_pdf[tar_list[tar_id][i].split('_')[2]] + '\n'
                ip_name_2 = 'Пакет: ' + ip_service_names_pdf[tar_list[tar_id][i].split('_')[3].split('-')[0]]
                ip_name = ip_name_1 + ip_name_2

                doc.tables[1].cell(1, i + 1).text = str(ip_name)
                doc.tables[1].cell(1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # iptv price
                ip_price = tar_list[tar_id][i].split('_')[3].split('-')[1] + ' ₽/мес.'

                doc.tables[1].cell(2, i + 1).text = str(ip_price)
                doc.tables[1].cell(2, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # dop_ch
                ih_dop = str()
                for j in tar_list[tar_id][i].split('_'):
                    try:
                        try:
                            ih_dop += ip_dop_ch_names_pdf[j.split('-')[0]] + f" ({ip_dop_ch_names_pdf[j.split('-')[1]]}) - {j.split('-')[2]} ₽" + '\n'
                        except:
                            ih_dop += ip_dop_ch_names_pdf[j.split('-')[0]] + f" ({j.split('-')[1]} ₽) " + '\n'
                    except:
                        pass
                if len(ih_dop) == 0:
                    ih_dop = '--'
                else:
                    ih_dop = ih_dop[:-2]

                doc.tables[1].cell(3, i + 1).text = str(ih_dop)
                doc.tables[1].cell(3, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                #videomuz
                ip_videomuz = '--' if 'videomuz' not in tar_list[tar_id][i] else \
                    str([ip_prokat_names_pdf[int(a.split('-')[1])] for a in tar_list[tar_id][i].split('_') if 'videomuz' in a][0])

                doc.tables[1].cell(4, i + 1).text = str(ip_videomuz)
                doc.tables[1].cell(4, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(4, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # content_contr
                ip_content_contr = '--' if 'content/contr' not in tar_list[tar_id][i] else '450 ₽'

                doc.tables[1].cell(5, i + 1).text = str(ip_content_contr)
                doc.tables[1].cell(5, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(5, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # adv/konst
                ip_videomuz = '--' if 'adv/konst' not in tar_list[tar_id][i] else '220 ₽'

                doc.tables[1].cell(6, i + 1).text = str(ip_videomuz)
                doc.tables[1].cell(6, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(6, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(6, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # iptv block

                ip_block = '--' if 'iptv/block' not in tar_list[tar_id][i] else '20 ₽'

                doc.tables[1].cell(7, i + 1).text = str(ip_block)
                doc.tables[1].cell(7, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(7, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(7, i + 1).paragraphs[0].runs[0].font.name = 'Arial'


            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        elif tar_order[tar_id] == 'internet':
            doc = Document(curdir.joinpath(f'{tar_order[tar_id].replace("/", "_")}_kp{len(tar_list[tar_id])}.docx'))
            for i in range(len(tar_list[tar_id])):

                # naimenovanie
                doc.tables[1].cell(1, i + 1).text = internet_tariffs_names_pdf[tar_list[tar_id][i].split('_')[5]]
                doc.tables[1].cell(1, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(1, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                # speed
                doc.tables[1].cell(2, i + 1).text = str(tar_list[tar_id][i].split('_')[4].replace('M', ' Мбит/с'))
                doc.tables[1].cell(2, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(2, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

                #price
                price = await getprice(tar_list[tar_id][i])
                price = int(price)
                sale = int(tar_list[tar_id][i].split('_')[6][:-1])
                if sale != 0:
                    price = price - (price * (sale / 100))
                # doc.tables[1].cell(3, i + 1).text = f'{await tostr_(list(str(int(price))))} ₽'
                doc.tables[1].cell(3, i + 1).text = str(int(price)) + ' ₽'
                doc.tables[1].cell(3, i + 1).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.size = Pt(9)
                doc.tables[1].cell(3, i + 1).paragraphs[0].runs[0].font.name = 'Arial'

            # Имя клиент
            doc.paragraphs[1].text = f"Уважаемый {tar_list[-1][1]} !"
            doc.paragraphs[1].runs[0].font.name = 'Arial'
            doc.paragraphs[1].runs[0].font.size = Pt(10)

            # Адрес
            doc.paragraphs[3].text = f"{tar_list[-1][0]}"
            doc.paragraphs[3].runs[0].font.name = 'Arial'
            doc.paragraphs[3].runs[0].font.size = Pt(9)
            doc.paragraphs[3].runs[0].font.color.rgb = RGBColor.from_string('FF3900')

            # имя опер
            doc.paragraphs[8].text += f"{tar_list[-1][2]}"
            doc.paragraphs[8].runs[0].font.name = 'Arial'
            doc.paragraphs[8].runs[0].font.size = Pt(9)

            # номер
            doc.paragraphs[10].text += f"{tar_list[-1][3]}"
            doc.paragraphs[10].runs[0].font.name = 'Arial'
            doc.paragraphs[10].runs[0].font.size = Pt(9)

            # цена 1
            if 'ропустит' in tar_list[-1][4]:
                doc.paragraphs[5].text += "рассчитывается индивидуально"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[5].text += f"{tar_list[-1][4]} ₽"
                doc.paragraphs[5].runs[0].font.name = 'Arial'
                doc.paragraphs[5].runs[0].font.size = Pt(6.5)
                doc.paragraphs[5].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            # цена 2
            if 'ропустит' in tar_list[-1][5]:
                doc.paragraphs[6].text += "рассчитывается индивидуально"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')
            else:
                doc.paragraphs[6].text += f"{tar_list[-1][5]} ₽"
                doc.paragraphs[6].runs[0].font.name = 'Arial'
                doc.paragraphs[6].runs[0].font.size = Pt(6.5)
                doc.paragraphs[6].runs[0].font.color.rgb = RGBColor.from_string('9A9A9A')

            doc.save(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ДОК')
            try:
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            except:
                time.sleep(1)
                convert(curdir/f'pdf_temp/{uid}_{tar_id}_kp_new.docx')
            print('СОЗДАН ПДФ')

        else:
            print('pochemu tak')
